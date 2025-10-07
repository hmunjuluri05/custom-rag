import PyPDF2
from docx import Document
import openpyxl
import pandas as pd
from pathlib import Path
from typing import Dict, Any, List
import logging

from .interfaces.document_processor_interface import IDocumentProcessor

logger = logging.getLogger(__name__)

class DocumentProcessor(IDocumentProcessor):
    """Handles text extraction from various document formats"""

    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_xlsx,
            '.xls': self._extract_excel,
            '.txt': self._extract_text
        }

    def extract_text(self, file_path: Path) -> str:
        """Extract text content from a document file"""

        file_extension = file_path.suffix.lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        try:
            extractor = self.supported_formats[file_extension]
            text_content = extractor(file_path)

            if not text_content.strip():
                logger.warning(f"No text content extracted from {file_path}")
                return ""

            return text_content

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from {file_path.name}: {str(e)}")

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        text_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1} from {file_path}: {str(e)}")
                    continue

        return "\n\n".join(text_content)

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        text_content = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))

            if table_text:
                text_content.append("--- Table ---\n" + "\n".join(table_text))

        return "\n\n".join(text_content)

    def _extract_xlsx(self, file_path: Path) -> str:
        """Extract text from XLSX files with enhanced multiple sheet support"""
        workbook = openpyxl.load_workbook(file_path)
        text_content = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_content = [f"=== SHEET: {sheet_name} ==="]

            # Get headers from first row
            headers = []
            first_row = True

            # Get all rows with data
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                row_data = [str(cell) if cell is not None else "" for cell in row]

                # Only include rows that have at least one non-empty cell
                if any(cell.strip() for cell in row_data if isinstance(cell, str)):
                    # First row is treated as headers
                    if first_row:
                        headers = row_data
                        sheet_content.append("COLUMNS: " + " | ".join(row_data))
                        first_row = False
                    else:
                        # Create structured row with column names
                        row_str = " | ".join(row_data)

                        # Also create a searchable format with column=value pairs
                        if headers:
                            structured_data = []
                            for col_name, value in zip(headers, row_data):
                                if value.strip():
                                    structured_data.append(f"{col_name}={value}")
                            if structured_data:
                                row_str += " [" + ", ".join(structured_data) + "]"

                        sheet_content.append(row_str)

            if len(sheet_content) > 1:  # More than just the header
                text_content.append("\n".join(sheet_content))

        return "\n\n".join(text_content)

    def _extract_excel(self, file_path: Path) -> str:
        """Extract text from XLS files using pandas with enhanced multiple sheet support"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                if not df.empty:
                    sheet_content = [f"=== SHEET: {sheet_name} ==="]

                    # Convert DataFrame to string representation
                    # Replace NaN values with empty strings
                    df_filled = df.fillna('')

                    # Add column headers
                    headers = [str(col) for col in df_filled.columns]
                    sheet_content.append("COLUMNS: " + " | ".join(headers))

                    # Add data rows with structured format
                    for _, row in df_filled.iterrows():
                        row_text = " | ".join(str(val) for val in row.values)

                        # Also create a searchable format with column=value pairs
                        structured_data = []
                        for col_name, value in zip(headers, row.values):
                            value_str = str(value).strip()
                            if value_str:
                                structured_data.append(f"{col_name}={value_str}")
                        if structured_data:
                            row_text += " [" + ", ".join(structured_data) + "]"

                        sheet_content.append(row_text)

                    text_content.append("\n".join(sheet_content))

            return "\n\n".join(text_content)

        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {str(e)}")
            raise Exception(f"Failed to read Excel file: {str(e)}")

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise Exception(f"Failed to read text file with UTF-8 or Latin-1 encoding: {str(e)}")

    def get_document_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from document"""

        stat_info = file_path.stat()

        metadata = {
            "filename": file_path.name,
            "file_extension": file_path.suffix.lower(),
            "file_size": stat_info.st_size,
            "created_time": stat_info.st_ctime,
            "modified_time": stat_info.st_mtime
        }

        # Add format-specific metadata
        file_extension = file_path.suffix.lower()

        if file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(pdf_reader.pages)
                    if pdf_reader.metadata:
                        metadata['pdf_metadata'] = {
                            'title': pdf_reader.metadata.get('/Title', ''),
                            'author': pdf_reader.metadata.get('/Author', ''),
                            'subject': pdf_reader.metadata.get('/Subject', ''),
                            'creator': pdf_reader.metadata.get('/Creator', '')
                        }
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata from {file_path}: {str(e)}")

        elif file_extension == '.xlsx':
            try:
                workbook = openpyxl.load_workbook(file_path)
                metadata['sheet_count'] = len(workbook.sheetnames)
                metadata['sheet_names'] = workbook.sheetnames
            except Exception as e:
                logger.warning(f"Could not extract Excel metadata from {file_path}: {str(e)}")

        return metadata

    # Interface implementation methods
    async def process_file(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Process a single file and return processed content"""
        try:
            # Extract text content
            text_content = self.extract_text(file_path)

            # Extract metadata
            metadata = self.get_document_metadata(file_path)

            # Get chunking parameters
            from ..embedding.chunking import ChunkerFactory, ChunkingConfig, ChunkingStrategy

            chunk_size = kwargs.get('chunk_size', 1000)
            chunk_overlap = kwargs.get('chunk_overlap', 200)
            strategy = kwargs.get('chunking_strategy', 'recursive_character')

            # Create chunking config
            chunking_config = ChunkingConfig(
                strategy=ChunkingStrategy(strategy),
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )

            # Create chunker and split text
            chunker = ChunkerFactory.create_chunker(chunking_config)
            chunks = await chunker.split_text(text_content)

            return {
                'chunks': chunks,
                'metadata': metadata,
                'original_text': text_content,
                'chunk_count': len(chunks)
            }

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise

    async def process_files(self, file_paths: List[Path], **kwargs) -> List[Dict[str, Any]]:
        """Process multiple files and return processed contents"""
        results = []
        for file_path in file_paths:
            try:
                result = await self.process_file(file_path, **kwargs)
                result['file_path'] = str(file_path)
                result['status'] = 'success'
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {str(e)}")
                results.append({
                    'file_path': str(file_path),
                    'status': 'error',
                    'error': str(e)
                })
        return results

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())

    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from a file (interface implementation)"""
        return self.get_document_metadata(file_path)